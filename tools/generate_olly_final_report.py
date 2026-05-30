from __future__ import annotations

import json
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "final_report_assets"
OUTPUT = DOCS / "OLLY_최종_보고서.docx"
DEMO_OBSERVATIONS = ASSETS / "demo_observations.json"

FONT_KO = "맑은 고딕"
FONT_LATIN = "Arial"
FONT_MONO = "Consolas"


def run_text(command: list[str], timeout: int = 8) -> str:
    try:
        result = subprocess.run(
            command,
            cwd=ROOT,
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=timeout,
            check=False,
        )
    except Exception as exc:
        return f"(명령 실행 실패: {exc})"
    return result.stdout.strip() or "(출력 없음)"


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, 9, bold, color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 24, "111827"),
        ("Heading 1", 17, "111827"),
        ("Heading 2", 14, "1F2937"),
        ("Heading 3", 12, "374151"),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
        code.font.name = FONT_MONO
        code._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO)
        code.font.size = Pt(8.5)

    if "Report Caption" not in styles:
        caption = styles.add_style("Report Caption", 1)
        caption.font.name = FONT_LATIN
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        caption.font.size = Pt(8.5)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string("4B5563")


def add_paragraph(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, 10, True)
        tail = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(tail, 10)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 10)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, 10)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, 10)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    cell.text = ""
    for idx, line in enumerate(text.strip().splitlines()):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.style = "Code Block"
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = FONT_MONO
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO)
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Report Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, 8, False, "4B5563")


def add_image(doc: Document, filename: str, caption: str, width_inches: float = 6.5) -> None:
    path = ASSETS / filename
    if not path.exists():
        add_code_block(doc, f"[캡쳐 파일 없음] {path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], "E5E7EB")
        set_cell_text(header_cells[idx], header, bold=True)
        if widths:
            header_cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def add_toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    note = doc.add_paragraph()
    note_run = note.add_run("※ Word에서 문서를 열고 목차를 선택한 뒤 '필드 업데이트'를 누르면 페이지 번호가 자동 갱신됩니다.")
    set_run_font(note_run, 9, False, "6B7280")


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(110)
    title_run = title.add_run("OLLY 최종 보고서")
    set_run_font(title_run, 26, True, "111827")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("실무형 LLM 서비스 관측성 플랫폼 MVP")
    set_run_font(subtitle_run, 15, True, "4F46E5")

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_before = Pt(20)
    run = summary.add_run("Kubernetes, OpenTelemetry, Prometheus, Jaeger 기반 운영 관측성 구현")
    set_run_font(run, 11, False, "374151")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(90)
    meta_run = meta.add_run("오픈소스분석(클라우드) | 5조")
    set_run_font(meta_run, 11, False, "111827")
    doc.add_page_break()


def add_overview_box(doc: Document) -> None:
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["프로젝트명", "OLLY"],
            ["핵심 목표", "LLM 요청의 비용, 토큰, 지연, 병목, 실패, 알림 원인을 데이터 기반으로 관측"],
            ["현재 MVP 대상", "Ollama gemma3:1b 기반 로컬 SLM 샘플 서비스"],
            ["실무 적용 방향", "사용자용 LLM/RAG 챗봇 서비스의 /chat API에 OLLY 계측을 연결"],
            ["실무성 기준", "기존 서비스 교체 없이 계측 추가, 표준 도구 기반 운영, 알림과 추적 자동화, 프롬프트 보안 고려"],
            ["CNCF 적용", "Kubernetes, OpenTelemetry, Prometheus, Jaeger"],
        ],
        [1.6, 4.8],
    )


def load_demo_observations() -> list[dict[str, Any]]:
    if not DEMO_OBSERVATIONS.exists():
        return []
    try:
        data = json.loads(DEMO_OBSERVATIONS.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    if not isinstance(data, list):
        return []
    return [item for item in data if isinstance(item, dict)]


def _format_demo_observation_row(item: dict[str, Any]) -> list[str]:
    scenario = str(item.get("scenario") or "-")
    request_id = str(item.get("request_id") or "-")
    trace_id = str(item.get("trace_id") or "")
    trace_short = trace_id[:12] if trace_id else "-"
    latency_ms = item.get("latency_ms")
    tokens = item.get("tokens")
    cost_usd = item.get("cost_usd")
    status = str(item.get("status") or "-")
    dominant_stage = str(item.get("dominant_stage") or "-")
    dominant_stage_ms = item.get("dominant_stage_ms")

    latency_text = f"{latency_ms}ms" if latency_ms is not None else "-"
    tokens_text = f"{tokens} tokens" if tokens is not None else "-"
    if isinstance(cost_usd, int | float):
        cost_text = f"${cost_usd:.8f}"
    else:
        cost_text = "-"
    if dominant_stage_ms is not None:
        stage_text = f"{dominant_stage} {dominant_stage_ms}ms"
    else:
        stage_text = dominant_stage

    return [
        scenario,
        f"{request_id}\ntrace {trace_short}",
        f"{status}\n{latency_text}",
        f"{tokens_text}\n{cost_text}",
        stage_text,
    ]


def add_demo_observation_table(doc: Document) -> None:
    observations = load_demo_observations()
    if not observations:
        add_paragraph(
            doc,
            "재검증 관측값 파일이 없는 경우에는 개별 측정값보다 span과 metric이 어떤 방향으로 증가하는지 중심으로 해석한다.",
        )
        return

    scenario_order = {"normal": 0, "slow_retrieve": 1, "slow_llm": 2, "high_token": 3, "error": 4}
    report_observations = [item for item in observations if str(item.get("scenario")) != "normal"] or observations
    rows = [
        _format_demo_observation_row(item)
        for item in sorted(report_observations, key=lambda row: scenario_order.get(str(row.get("scenario")), 99))
    ]
    add_table(
        doc,
        ["시나리오", "요청 식별자", "상태/지연", "토큰/비용", "주요 관측 단계"],
        rows,
        [1.05, 1.35, 1.2, 1.35, 1.7],
    )


def main() -> None:
    doc = Document()
    configure_document(doc)

    health = run_text(["curl", "-sS", "--max-time", "5", "http://localhost:8001/health"])
    k8s_status = run_text(["kubectl", "-n", "olly", "get", "pods,svc,pvc,job"], timeout=12)

    add_title_page(doc)
    doc.add_heading("문서 목차", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("1. 서론", level=1)
    add_paragraph(
        doc,
        "LLM 기반 서비스는 일반적인 웹 API보다 운영 원인 분석이 복잡하다. 동일한 /chat API 호출 시에도 "
        "프롬프트 길이, 검색 문맥, 모델 종류 및 출력 길이에 따라 시스템 리소스 비용과 응답 시간의 편차가 크게 발생한다. 사용자에게는 "
        "답변 지연 또는 실패라는 결과만 드러나지만, 운영자는 검색 단계, 모델 호출 단계, 후처리 단계 중 어느 단계에서 "
        "문제가 발생하였는지 별도로 추적해야 한다.",
    )
    add_paragraph(
        doc,
        "OLLY는 이러한 운영 문제를 해결하기 위해 LLM 요청을 request_id와 trace_id 중심으로 연결하고, 비용, 토큰, "
        "latency, 단계별 병목, 실패, 알림을 통합적으로 제공하는 관측성 MVP이다. 본 프로젝트에서는 "
        "로컬 SLM 기반 샘플 서비스를 대상으로 관측성 구조를 구현하고 실제 LLM/RAG 챗봇 서비스에 적용 가능한 형태로 확장 가능성을 검증하였다. 특히 실무 환경에서 요구되는 "
        "표준 계측, 장애 추적, 비용 관리, 알림, 보안 고려를 MVP 범위 안에서 검증하도록 설계하였다.",
    )
    add_overview_box(doc)

    doc.add_heading("2. 프로젝트 목표와 적용 범위", level=1)
    doc.add_heading("2.1 관측 대상의 정의", level=2)
    add_paragraph(
        doc,
        "본 보고서에서 LLM 서비스는 사용자의 자연어 요청을 받아 모델 호출 또는 RAG 검색을 거쳐 답변을 반환하는 서비스를 의미한다. "
        "Claude, Gemini, OpenAI 같은 외부 LLM API를 사용하는 서비스, 내부 문서 검색을 결합한 RAG 챗봇, Ollama 기반 "
        "로컬 SLM 서비스가 모두 이 범주에 포함된다.",
    )
    add_bullets(
        doc,
        [
            "외부 LLM API형 서비스: 모델 제공사의 API를 호출하고 토큰 기반 비용이 발생한다.",
            "RAG 기반 챗봇형 서비스: 검색 단계와 모델 생성 단계가 분리되므로 단계별 병목 분석이 중요하다.",
            "로컬 SLM형 서비스: API 과금은 없지만 CPU/GPU 실행 시간과 인프라 비용이 주요 관측 대상이 된다.",
        ],
    )
    doc.add_heading("2.2 MVP 구현 범위", level=2)
    add_paragraph(
        doc,
        "현재 구현은 비용과 재현성을 고려하여 Ollama gemma3:1b를 사용하는 FastAPI 샘플 서비스를 관측 대상으로 삼았다. "
        "사용자 챗봇 UI는 이 샘플 서비스에 질문을 보내고, 운영자 대시보드와 분석 챗봇은 Prometheus와 Jaeger에 쌓인 "
        "관측 데이터를 조회한다. 따라서 현재 OLLY는 로컬 SLM 기반 샘플 서비스를 관측하는 독립 실행형 MVP이며, "
        "운영자 분석 챗봇은 관측 결과를 해석하는 인터페이스로 정의된다.",
    )
    doc.add_heading("2.3 실무 적용 방향", level=2)
    add_paragraph(
        doc,
        "실무에서 OLLY는 사용자용 챗봇 서비스를 대체하지 않고, 기존 LLM/RAG 챗봇 API에 연결되는 관측성 계층으로 동작한다. "
        "서비스 코드에 OpenTelemetry 계측과 metric 기록을 추가하면 OLLY 대시보드가 해당 요청의 request_id, trace_id, "
        "모델, 토큰, 비용, latency, 단계별 duration, alert를 조회한다.",
    )
    add_code_block(
        doc,
        """
사용자용 LLM/RAG 챗봇
  -> /chat API
  -> retrieve 또는 tool call
  -> llm_call
  -> postprocess

OLLY 관측 계층
  -> OpenTelemetry trace
  -> Prometheus metric
  -> Jaeger 요청 추적 상세
  -> 운영자 대시보드 / 분석 챗봇
        """,
    )
    doc.add_heading("2.4 실무 적용성을 고려한 설계 원칙", level=2)
    add_paragraph(
        doc,
        "OLLY의 실무 적용성은 이미 운영 중인 사용자용 LLM/RAG 서비스에 관측성 계층을 유연하게 통합 적용할 수 있다는 점에 있다. "
        "따라서 구현에서는 특정 벤더의 전용 기능보다 OpenTelemetry, Prometheus, Jaeger처럼 "
        "운영 현장에서 널리 쓰이는 표준 도구와 인터페이스를 우선하였다.",
    )
    add_bullets(
        doc,
        [
            "기존 서비스 침투 최소화: /chat API의 핵심 로직은 유지하고 trace와 metric 기록을 추가하는 구조로 설계하였다.",
            "운영자 워크플로우 반영: KPI 확인, request_id 조회, trace 분석, alert 확인, Discord 알림까지 장애 대응 흐름을 연결하였다.",
            "벤더 종속성 완화: 외부 LLM API, RAG 챗봇, 로컬 SLM 모두 같은 관측 모델로 수집하도록 설계하였다.",
            "보안과 개인정보 고려: raw prompt 저장보다 template id, token count, 익명화된 식별자 중심의 관측을 권장한다.",
            "확장 가능한 배포 단위: Docker Compose와 Kubernetes manifest를 모두 제공해 로컬 검증과 클러스터 배포를 구분하였다.",
        ],
    )
    doc.add_heading("2.5 제약 조건을 반영한 설계 의사결정", level=2)
    add_paragraph(
        doc,
        "OLLY의 최종 아키텍처는 비용, 기간, 재현 가능성 및 운영 흐름 등의 현실적 제약 조건을 종합적으로 고려하여 도출된 최적화 결과물이다. "
        "이 절에서는 최종 결과물에 직접 영향을 준 핵심 의사결정을 정리한다.",
    )
    add_paragraph(
        doc,
        "웹 기반 통합 인터페이스 구현: 요청 생성, 메타데이터 확인, 추적 분석을 동일한 운영 흐름 안에서 연결하기 위해 Chat UI, Dashboard, 분석 챗봇을 포함한 웹 서비스 형태로 구현하였다. "
        "사용자가 질문을 보내면 request_id, trace_id, latency, token, cost가 생성되고, 운영자는 같은 요청을 대시보드와 Jaeger에서 연계하여 확인한다. "
        "이 구성은 운영자가 실제로 확인하는 화면과 장애 대응 흐름을 포함하므로 실무 적용성을 높인다.",
        bold_prefix="웹 기반 통합 인터페이스 구현:",
    )
    add_paragraph(
        doc,
        "로컬 SLM 기반 관측 대상 선정: Claude, Gemini, OpenAI 같은 외부 LLM API를 직접 연결하면 반복 검증 과정에서 비용과 호출량 제한 문제가 발생할 수 있다. "
        "본 프로젝트에서는 Ollama gemma3:1b 기반 로컬 SLM을 MVP 관측 대상으로 선택하였다. 이 결정으로 외부 API 비용 없이 반복 검증이 가능해졌고, "
        "로컬 모델 실행 시간을 기반으로 local infra cost를 추정하는 metric을 구현하였다. 이러한 접근은 비용 관측이라는 핵심 목표를 유지함과 동시에, 제한된 실험 환경에서의 리소스 제약 및 시스템 검증의 안정성 문제를 효과적으로 해결한다.",
        bold_prefix="로컬 SLM 기반 관측 대상 선정:",
    )
    add_paragraph(
        doc,
        "관측 대상 샘플 챗봇까지 함께 구현: request_id, trace_id, token, cost는 실제 요청 처리 과정에서 생성되어야 관측 대상으로서 의미를 갖는다. "
        "본 프로젝트에서는 OLLY 대시보드 구축과 함께 관측 대상이 되는 샘플 /chat API와 Chat UI를 통합하여 구현하였다. 그 결과 채팅 요청 하나마다 비용, 토큰, latency, trace가 생성되고, "
        "normal, slow_retrieve, slow_llm, high_token, error 시나리오로 관측 결과를 검증한다. 이를 통해 관측 플랫폼의 효용성을 입증하기 위한 샘플 서비스 환경을 통합적으로 설계 및 구축하였다.",
        bold_prefix="관측 대상 샘플 챗봇까지 함께 구현:",
    )
    add_paragraph(
        doc,
        "Discord webhook 기반 외부 알림 연계: 웹 화면 내부 알림은 실제 운영 환경의 장애 전파 채널로 한계가 있다. "
        "본 프로젝트에서는 Prometheus alert와 커스텀 alert rule을 Discord webhook과 연결해 외부 알림 채널까지 포함하였다. "
        "임계값 초과 시 Discord로 알림을 전송하고, 선택적으로 LLM 한 줄 요약을 첨부하여 운영자가 이상 신호의 의미를 빠르게 파악하도록 설계하였다. "
        "장애 신호가 대시보드 안에 머무르지 않고 운영 커뮤니케이션 채널로 전달되는 흐름을 반영하여 실제 운영 환경의 알림 흐름에 부합한다.",
        bold_prefix="Discord webhook 기반 외부 알림 연계:",
    )
    doc.add_heading("2.6 운영 관측성 중심의 평가 범위", level=2)
    add_paragraph(
        doc,
        "본 MVP의 관측 범위는 LLM 서비스의 운영 안정성과 인프라 리소스 효율성(Operational Observability)으로 명확히 한정한다. "
        "모델의 언어적 성능, 답변 품질, 할루시네이션과 같은 Semantic Monitoring 영역은 향후 확장 과제로 분리하며, "
        "본 프로젝트에서는 시스템 성능(Latency, Cost, Token, Error)을 추적하기 위한 데이터 파이프라인과 분산 추적(Distributed Tracing) 아키텍처 구축에 집중한다.",
    )
    add_paragraph(
        doc,
        "평가 단위는 하나의 /chat 요청이며, request_id, trace_id, token, cost, latency, stage duration을 공통 관측 변수로 사용한다. "
        "반복 가능한 검증을 위해 로컬 SLM을 사용하고, normal, slow_retrieve, slow_llm, high_token, error 시나리오로 병목과 실패 유형을 통제하였다.",
    )
    add_paragraph(
        doc,
        "생성형 AI는 코드 초안, 배포 설정 초안, 문서 초안 작성에 제한적으로 활용되었으나, 평가 변수, 기능 범위, 배포 방식, 검증 시나리오는 "
        "위 관측성 기준에 따라 확정하였다.",
    )
    doc.add_heading("2.7 관측 단위와 데이터 기준", level=2)
    add_paragraph(
        doc,
        "앞선 의사결정에 따라 OLLY의 관측 범위는 운영자가 실제로 판단할 수 있는 데이터에 맞추어 정리되었다. 본 프로젝트는 운영자가 비용과 성능 문제를 재현 가능한 데이터로 "
        "추적하게 하는 운영 관측성(Operational Observability)에 초점을 맞추었다.",
    )
    add_paragraph(
        doc,
        "프롬프트 관측은 서비스 코드에서 남길 정보를 결정해야 하는 계측 설계 문제로 정의하였다. 동일한 /chat API라도 request_id, trace_id, feature, prompt_template_id, token_count, latency를 "
        "함께 남기면 요청별 차이를 명확히 비교할 수 있다.",
    )
    add_paragraph(
        doc,
        "특히 실제 서비스 운영 시 발생할 수 있는 보안 및 프라이버시 문제를 고려하여, 사용자의 원시 프롬프트(Raw Prompt) 전체를 데이터베이스에 직접 저장하는 방식을 지양하였다. "
        "대신 prompt_template_id, 프롬프트 길이, 토큰 수, 그리고 익명화된 식별자(User/Org ID) 중심의 메타데이터만 수집하도록 데이터 기준을 확립하여 실무 적용의 안정성을 높였다.",
    )

    doc.add_heading("3. 시스템 아키텍처", level=1)
    add_paragraph(
        doc,
        "OLLY의 전체 시스템 아키텍처는 세 가지 독립적인 파이프라인(Pipeline)으로 구조화되었다. 첫째는 사용자가 질문을 보내고 답변을 받는 요청 처리 경로, 둘째는 "
        "애플리케이션에서 생성된 trace와 metric이 수집되는 관측성 경로, 셋째는 운영자가 대시보드와 분석 챗봇을 통해 원인을 찾는 운영 분석 경로이다. "
        "이 분리는 실제 서비스에 OLLY를 적용할 때 업무 로직과 관측 파이프라인을 독립적으로 확장하기 위한 구조적 기준이다.",
    )
    doc.add_heading("3.1 요청 처리 흐름", level=2)
    add_paragraph(
        doc,
        "사용자가 Chat UI를 통해 질의를 입력하면, 해당 요청은 FastAPI 기반의 /chat 엔드포인트로 인입된다. 이후 시스템은 내부적으로 "
        "문맥을 구성하는 검색(retrieve) 단계, Ollama(gemma3:1b)를 통해 실질적인 답변을 생성하는 모델 추론(llm_call) 단계, "
        "그리고 최종 응답을 정제하는 후처리(postprocess) 단계를 순차적으로 거쳐 사용자에게 최종 답변을 반환하도록 설계되었다.",
    )
    doc.add_heading("3.2 관측성 데이터 흐름", level=2)
    add_paragraph(
        doc,
        "FastAPI 애플리케이션 내부에서 발생한 계측 데이터는 OpenTelemetry SDK를 통해 수집기(OpenTelemetry Collector)로 스트리밍된다. "
        "수집된 데이터 중 요청 수, 토큰 사용량, 비용, 응답 지연(latency), 에러 및 단계별 소요 시간(stage duration)과 같은 정량적 시계열 지표는 "
        "Prometheus로 라우팅되어 저장된다. 동시에 POST /chat 요청의 전체 트레이스(Trace)와 내부 3단계(retrieve, llm_call, postprocess)의 "
        "하위 스팬(Span) 정보는 분산 추적 시스템인 Jaeger로 전달되어 상세한 데이터 추적 기반을 구성한다.",
    )
    doc.add_heading("3.3 운영 분석 흐름", level=2)
    add_paragraph(
        doc,
        "운영자가 분석 챗봇에 시스템 상태에 대한 질의를 입력하면, analysis_intents.py 모듈이 질문의 의도와 조회 기간, 그리고 특정 request_id 및 trace_id를 파싱하여 식별한다. "
        "이후 analysis.py 모듈이 전체 오케스트레이션을 담당하여 Prometheus와 Jaeger의 API를 직접 호출해 필요한 데이터를 조회한다. "
        "최종적으로 조회된 정량적 수치 근거를 바탕으로 시스템의 현재 상태나 병목 원인을 분석한 한국어 운영 답변이 운영자에게 반환되는 구조이다.",
    )
    add_image(doc, "15_architecture_overview.png", "그림 1. OLLY 전체 아키텍처: 사용자 요청 처리, 관측성 수집, 운영 분석과 알림 흐름", 6.8)
    add_image(doc, "16_user_flow.png", "그림 2. OLLY 요청 처리 및 운영자 분석 플로우: request_id와 trace_id 기반 추적 흐름", 6.8)
    add_paragraph(
        doc,
        "전체 구조는 운영 환경에서 필요한 관심사 분리(Separation of Concerns)를 기준으로 구성하였다. 사용자 요청 처리 경로와 관측 데이터 수집 경로를 엄격히 분리하고, "
        "운영자 화면은 Prometheus와 Jaeger에 저장된 데이터를 수동적으로 조회하는 소비자로 두었다. 이러한 방식은 실제 프로덕션 환경에 OLLY를 이식할 때에도 "
        "애플리케이션 코드, 관측 파이프라인, 운영 화면을 상호 의존성 없이 독립적으로 교체하거나 확장할 수 있는 유연한 아키텍처를 제공한다.",
    )
    add_image(doc, "01_chat_ui.png", "그림 3. 사용자/운영자 챗봇 UI: 질문, 응답, request metadata를 확인하는 화면")
    add_image(doc, "02_dashboard.png", "그림 4. OLLY 운영자 대시보드: KPI, 최근 요청, 병목 단계, 알림 상태를 통합 표시")

    doc.add_heading("4. 핵심 기능 구현", level=1)
    add_paragraph(
        doc,
        "OLLY의 핵심 기능은 요청 단위 추적, 운영 지표 분석, 알림 기반 대응으로 요약된다. 사용자가 보낸 하나의 질문은 request_id와 trace_id를 기준으로 "
        "대시보드, Jaeger trace, Prometheus metric에서 동일하게 추적된다. 운영자는 이 연결을 통해 평균 응답 시간과 함께 "
        "특정 요청이 어느 단계에서 비용과 토큰 사용량을 발생시켰는지 정밀하게 추적한다.",
    )
    add_paragraph(
        doc,
        "본 장은 운영자가 실제로 수행하는 판단 과정을 중심으로 핵심 기능을 설명한다. 질문 생성, 지표 확인, trace 분석, 알림 확인이 "
        "하나의 흐름으로 연결되어야 OLLY의 관측성 기능이 의미를 갖기 때문이다.",
    )
    doc.add_heading("4.1 사용자 챗봇 UI", level=2)
    add_paragraph(
        doc,
        "사용자는 /chat-ui에서 질문을 입력한다. 응답에는 answer와 함께 request_id, trace_id, latency, input/output token, "
        "cost, model, status가 함께 표시된다. 이 metadata가 운영자 대시보드와 Jaeger trace를 연결하는 핵심 식별자이다.",
    )
    add_paragraph(
        doc,
        "Chat UI는 사용자를 위한 최종 서비스 프론트엔드임과 동시에, 관측 데이터가 최초로 생성되는 진입점(Entry Point) 역할을 수행한다. 같은 화면에서 빠른 질문과 시나리오 요청을 반복 실행할 수 있기 때문에, "
        "정상 요청, 지연, 토큰 증가, 실패 요청의 관측 데이터를 동일한 인터페이스에서 생성할 수 있다.",
    )
    add_paragraph(
        doc,
        "특히 본 프로젝트의 핵심 검증 목적은 관측 데이터의 생성 및 추적 가능성 확인에 있다. "
        "이를 위해 백엔드 모델 제어 모듈에서 Ollama API(/api/chat)를 직접 호출하여 temperature, top_p, num_predict 등의 하이퍼파라미터를 로우레벨에서 제어하도록 구현하였다. "
        "동시에 검증 과정에서 발생할 수 있는 모델 응답의 불확실성을 통제하기 위해, 시나리오 응답 제어 모듈 기반의 응답 안정화 레이어(Stabilization Layer)를 설계하였다. "
        "이를 통해 시나리오별 관측 메타데이터는 실제 요청 처리 경로와 동일하게 발생시키면서도 설명 문구는 일관되게 제공하는 테스트 통제 환경을 완성하였다.",
    )
    add_image(doc, "06_chat_quick_question_selected.png", "그림 5. 빠른 질문 선택: 자주 묻는 운영 질문을 클릭해 입력창에 바로 채우는 화면")
    add_image(doc, "07_chat_quick_status_summary.png", "그림 6. 빠른 질문 응답: 상태 요약 답변과 request_id, trace_id, latency, token, cost metadata")
    doc.add_heading("4.2 운영자 대시보드", level=2)
    add_paragraph(
        doc,
        "본 대시보드는 다중 데이터소스를 직접 통합하는 데이터 집계(Aggregation) API를 중심으로 구현되었다. "
        "대시보드 백엔드 모듈은 Prometheus에서 KPI 및 병목 구간 분석을 위한 스칼라/벡터 쿼리(PromQL)를 수행함과 동시에, Jaeger API를 호출하여 최근 Trace 데이터를 수집한다.",
    )
    add_paragraph(
        doc,
        "특히 Jaeger API의 JSON 응답을 그대로 화면에 표출하는 수준에 머무르지 않고, Trace 내부의 Root Span과 Stage Span을 순회하며 tag 데이터를 파싱하는 커스텀 알고리즘을 적용하였다. "
        "이를 통해 olly.request_id, olly.cost_usd, olly.input_tokens, olly.output_tokens, 단계별 소요 시간 등을 백엔드에서 직접 추출하고, "
        "프론트엔드가 요구하는 recent_requests, stage_bottleneck_summary, primary_insight 요약 모델 구조로 재구성하여 제공하였다.",
    )
    add_image(doc, "13_dashboard_traces.png", "그림 7. 요청 추적 화면: request_id와 trace_id 기준으로 최근 요청과 단계별 병목을 확인")
    add_image(doc, "12_dashboard_signals.png", "그림 8. 실시간 지표 화면: Grafana 패널로 p95 latency, token trend, cost, stage p95를 확인")
    doc.add_heading("4.3 분석 챗봇", level=2)
    add_paragraph(
        doc,
        "분석 챗봇은 자연어 의도 분류(Intent Classification) 모듈을 통해 사용자의 운영 질문을 비용, 지연, 병목, Trace 상세, error, alert 등의 구체적인 intent로 사전 라우팅(Routing)한다.",
    )
    add_paragraph(
        doc,
        "이후 운영 분석 코어 모듈이 분류된 intent를 바탕으로 Prometheus와 Jaeger에서 정량적 metric과 trace 데이터를 직접 조회한다. "
        "챗봇은 이렇게 확보된 수치적 근거(Grounded Data)를 바탕으로 최종 답변을 오케스트레이션(Orchestration)하여 생성함으로써, "
        "비정형 생성 모델에 대한 의존도를 낮추고 인프라 사실에 기반한 운영 분석을 제공하도록 구현되었다.",
    )
    doc.add_heading("4.4 비용, 토큰, 병목 분석", level=2)
    add_paragraph(
        doc,
        "OLLY는 요청별 비용과 토큰 사용량을 metric으로 저장하고, feature와 scenario 단위로 집계하도록 구현하였다. "
        "또한 retrieve, llm_call, postprocess span을 분리하여 전체 시스템의 응답 지연(Latency) 발생 시, 그 근본 원인(Root Cause)이 되는 단계를 명확히 식별한다. "
        "이 구조는 RAG 검색 지연과 모델 생성 지연을 구분하는 데 핵심적인 역할을 한다.",
    )
    add_paragraph(
        doc,
        "실무에서 LLM 비용은 총액과 기능별 요청 패턴의 관계를 함께 분석해야 한다. OLLY는 token, cost, latency, stage duration을 "
        "하나의 요청 단위로 묶어 특정 기능의 비용 증가가 실제 병목이나 실패와 연결되는지 함께 확인하도록 설계하였다.",
    )
    add_paragraph(
        doc,
        "현재 MVP는 외부 LLM API가 아니라 로컬 SLM인 Ollama gemma3:1b를 사용하므로 실제 토큰 API 과금은 발생하지 않는다. "
        "구현에서는 이를 반영하여 pricing.py에서 gemma3:1b의 input/output token 단가를 모두 0으로 설정하였다. "
        "대신 로컬 실행 비용은 추정 인프라 비용으로 모델링하였다. /chat 처리 시 Ollama metadata의 total duration 또는 llm_call 실행 시간을 compute_seconds로 기록하고, "
        "infra_cost_usd = compute_seconds / 3600 * LOCAL_COMPUTE_HOURLY_USD 공식을 적용한다. 본 검증 환경에서는 로컬 컴퓨팅 단위 비용(LOCAL_COMPUTE_HOURLY_USD)을 0.05 USD/hour로 책정하여, "
        "CPU 기반 로컬 추론 소요 시간을 정량적인 비용 지표로 환산하였다.",
    )
    add_paragraph(
        doc,
        "따라서 응답 metadata와 metric의 cost_usd는 실제 결제 금액이 아니라 token_cost_usd와 infra_cost_usd를 합산한 운영 관측용 추정치이다. "
        "로컬 SLM에서는 token_cost_usd가 0이며, 비용 변동은 주로 llm_call 실행 시간과 compute_seconds 증가에 의해 설명된다. "
        "이 구조는 향후 외부 LLM API로 전환할 경우 토큰 단가 테이블만 조정하여 동일한 cost metric 체계를 재사용할 수 있도록 하기 위한 설계이다.",
    )
    doc.add_heading("4.5 알림", level=2)
    add_paragraph(
        doc,
        "Prometheus alert rule로 p95 latency, error rate, token usage spike, retrieve slow, local inference cost spike를 감시한다. "
        "또한 운영자가 직접 metric, 조건, 임계값, 평가 window, cooldown, Discord webhook을 설정하는 커스텀 알림 기능을 구현하였다.",
    )
    add_paragraph(
        doc,
        "실무 환경에서는 장애 발생 후 관리자가 대시보드를 수동으로 모니터링하는 방식만으로는 신속한 초기 대응에 한계가 존재한다. OLLY는 임계값 기반 알림과 LLM 한 줄 요약을 "
        "함께 제공해 운영자가 문제 발생 지표와 원인을 빠르게 파악하도록 설계하였다.",
    )
    add_paragraph(
        doc,
        "또한 사용자가 생성한 커스텀 알림 규칙이 시스템 재시작 시 휘발되지 않도록 알림 영구 저장(Alert Storage) 모듈을 구현하여 로컬 파일 시스템에 상태를 유지(Persistence)하였다. "
        "비동기(Async) 처리 환경에서 여러 규칙이 동시에 조작될 때 발생할 수 있는 데이터 레이스(Data Race)를 방지하기 위해 asyncio.Lock을 적용하였다. "
        "나아가 임시 파일 교체(Atomic Replace) 방식을 적용함으로써, 규칙 저장 중 발생할 수 있는 데이터 손상 및 결함을 방지하는 운영 설정 지속성을 확보하였다.",
    )
    add_image(doc, "14_dashboard_alert_rules.png", "그림 9. 알림 관리 화면: 커스텀 알림 규칙, Discord webhook, 최근 발화 이력")

    doc.add_heading("5. CNCF 기반 관측성 구성", level=1)
    add_paragraph(
        doc,
        "본 프로젝트의 CNCF 기반 관측성 구성은 Kubernetes, OpenTelemetry, Prometheus, Jaeger를 중심으로 설계하였다. "
        "수집된 시계열 데이터의 직관적인 가시화를 위해 Grafana를 연동하여 모니터링 환경을 완성하였다.",
    )
    add_paragraph(
        doc,
        "CNCF 도구 조합은 LLM 요청의 운영 특성을 기준으로 선정하였다. LLM 기반 /chat 요청은 검색 단계, 모델 호출 단계, 후처리 단계가 결합되어 있어 "
        "일반 로그 수집은 병목 구간과 비용 증가 원인을 요청 단계별로 구조화하는 데 한계가 있다. 이에 따라 OpenTelemetry의 표준 Trace/Span 구조로 요청 내부 단계를 계측하고, "
        "Prometheus로 시계열 지표와 알림 조건을 관리하며, Jaeger로 특정 request_id와 trace_id의 요청 흐름을 시각적으로 검증하는 아키텍처를 채택하였다.",
    )
    add_paragraph(
        doc,
        "Kubernetes는 sample-llm-api, Ollama, OpenTelemetry Collector, Prometheus, Jaeger, Grafana를 컨테이너 단위로 실행하고 연결하는 기반으로 사용하였다. "
        "Deployment와 Service는 각 구성요소의 실행과 네트워크 접근을 정의하고, PVC는 모델 데이터 보존, Job은 gemma3:1b 모델 다운로드, ConfigMap은 collector와 Prometheus 설정을 관리한다. "
        "이를 통해 제3자도 동일한 manifest를 활용하여 관측 환경을 동일하게 재현할 수 있다.",
    )
    add_paragraph(
        doc,
        "OpenTelemetry는 FastAPI의 /chat 요청과 내부 단계인 retrieve, llm_call, postprocess를 span으로 나누는 계측 표준으로 적용하였다. "
        "특정 요청이 느릴 때 전체 API latency와 함께 검색 단계, 모델 호출 단계, 후처리 단계의 소요 시간을 trace 구조로 확인한다. "
        "이 계측은 서비스 코드의 표준 instrumentation에 기반하므로 외부 LLM API와 RAG 서비스에도 같은 방식으로 확장 가능하다.",
    )
    add_paragraph(
        doc,
        "Prometheus는 요청 수, latency, token, cost, error, stage duration을 metric으로 저장하고 PromQL 분석과 alert rule의 근거가 된다. "
        "운영자 대시보드와 분석 챗봇은 이 metric을 조회하여 기간별 요청량, 비용 증가, p95 latency, error rate, 토큰 급증 여부를 설명한다. "
        "Jaeger는 request_id와 trace_id를 기준으로 특정 요청의 span을 시각화하여 Prometheus의 집계 지표가 어떤 실제 요청에서 발생하였는지 확인하게 한다.",
    )
    add_image(doc, "17_observability_pipeline.png", "그림 10. CNCF 관측성 파이프라인: OpenTelemetry 수집, Prometheus metric, Jaeger trace 분석", 6.8)
    add_table(
        doc,
        ["CNCF 도구", "프로젝트 내 역할", "구현 근거", "효과"],
        [
            ["Kubernetes", "kind 기반 로컬 클러스터 실행", "Deployment, Service, PVC, Job, ConfigMap, NodePort", "서비스 단위 배포와 재현 가능한 실행 환경"],
            ["OpenTelemetry", "애플리케이션 계측 표준", "FastAPI, retrieve, llm_call, postprocess span 생성", "요청 흐름을 vendor-neutral 방식으로 추적"],
            ["Prometheus", "metric 저장과 PromQL 분석", "requests, tokens, cost, latency, errors, alert rules", "운영 지표 집계와 알림"],
            ["Jaeger", "분산 trace 시각화", "POST /chat trace와 단계별 span 확인", "요청 하나의 병목 위치를 확인"],
        ],
        [1.25, 1.7, 2.1, 1.7],
    )
    add_paragraph(
        doc,
        "이 조합은 실무 운영 환경에서도 설명 가능하다. OpenTelemetry는 애플리케이션 계측을 표준화하고, Prometheus는 metric과 alert "
        "운영을 담당하며, Jaeger는 요청 단위 원인 분석을 지원한다. OLLY가 특정 클라우드나 LLM 벤더의 전용 기능에 의존하지 않도록 "
        "구성한 이유도 실제 서비스에 적용할 때 교체 비용과 종속성을 줄이기 위해서이다.",
    )
    add_image(doc, "03_jaeger_trace.png", "그림 11. Jaeger trace: POST /chat 요청 안에서 retrieve, llm_call span을 확인")
    add_image(doc, "04_prometheus_alerts.png", "그림 12. Prometheus alerts: OLLY MVP alert rule 상태 확인")

    doc.add_heading("6. Docker 및 Kubernetes 배포", level=1)
    add_image(doc, "18_kubernetes_deployment.png", "그림 13. Kubernetes(kind) 배포 구조: namespace, Deployment, Service, PVC, Job, NodePort 연결", 6.8)
    doc.add_heading("6.1 Docker Compose", level=2)
    add_paragraph(
        doc,
        "Docker Compose는 로컬 데모용 실행 환경이다. sample-llm-api, ollama, otel-collector, jaeger, prometheus, grafana를 "
        "함께 실행하며, 첫 실행 시 ollama-pull-gemma 컨테이너가 gemma3:1b 모델을 내려받는다.",
    )
    add_code_block(
        doc,
        """
docker compose -f deploy/docker-compose.yml up -d --build
curl http://localhost:8001/health
        """,
    )
    doc.add_heading("6.2 Kubernetes(kind)", level=2)
    add_paragraph(
        doc,
        "Kubernetes 배포는 kind 기반 로컬 클러스터에서 검증하였다. Namespace, PVC, Deployment, Service, Job, ConfigMap을 사용하며, "
        "NodePort와 kind extraPortMappings를 통해 localhost에서 각 도구에 접근한다.",
    )
    add_paragraph(
        doc,
        "kind와 NodePort는 로컬 검증 및 실험 환경에서 시스템의 재현성을 확보하기 위한 아키텍처 구성이다. 실무 환경으로 확장할 때는 동일한 Deployment와 Service 구조를 "
        "기반으로 Ingress, TLS, Secret, 외부 저장소, 장기 metric 보관, 접근 권한 제어를 추가하는 방식으로 발전시킬 수 있다.",
    )
    add_code_block(doc, "cd k8s\nmake up\nmake status")
    add_paragraph(doc, "검증 시점의 /health 응답은 다음과 같다.")
    add_code_block(doc, health)
    add_paragraph(
        doc,
        "Kubernetes 상태 검증에서는 sample-llm-api, ollama, otel-collector, Prometheus, Jaeger, Grafana 리소스가 실행되고 "
        "ollama-pull-gemma Job이 완료되는지 확인하였다. 본문에는 배포 방식과 검증 의미를 중심으로 정리하고, 전체 kubectl 출력은 부록 C에 수록하였다.",
    )

    doc.add_heading("7. 검증 결과", level=1)
    add_paragraph(
        doc,
        "검증은 normal, slow_retrieve, slow_llm, high_token, error 시나리오를 기준으로 수행하였다. 각 검증 요청은 "
        "생성된 뒤 동일한 request_id와 trace_id로 대시보드, Jaeger, Prometheus 지표까지 연결되는 과정을 확인하는 방식으로 구성하였다. "
        "각 시나리오에서는 Chat UI 응답 metadata, Dashboard의 요청 상세, Jaeger span, Prometheus/Grafana 지표가 같은 요청을 "
        "설명하는지 확인하였다. 이 시나리오들은 운영 관측성 검증에 필요한 대표 사례이다.",
    )
    doc.add_heading("7.1 검증 시나리오 설계", level=2)
    add_paragraph(
        doc,
        "검증 시나리오는 실제 LLM 서비스 운영에서 자주 발생할 수 있는 문제 유형의 관측 신호를 분리하여 확인하도록 구성하였다. "
        "normal은 기준 상태, slow_retrieve는 RAG 검색 병목, slow_llm은 모델 생성 병목, high_token은 비용 증가, error는 실패 추적을 대표한다. "
        "이 다섯 가지는 운영자가 대시보드에서 원인을 판단해야 하는 대표 상황이다.",
    )
    add_paragraph(
        doc,
        "각 시나리오는 Chat UI에서 요청을 생성하고, 응답 metadata에 포함된 request_id와 trace_id를 기준으로 Dashboard, Jaeger, Prometheus 지표가 같은 요청을 설명하는지 확인하는 방식으로 검증하였다. "
        "검증의 핵심은 하나의 요청이 관측 데이터로 연결되고 운영자가 그 데이터를 근거로 병목과 비용 원인을 판단할 수 있음을 확인하는 데 있다.",
    )
    doc.add_heading("7.2 시나리오별 관측 결과 요약", level=2)
    add_paragraph(
        doc,
        "각 시나리오는 Chat UI에서 발생한 단일 요청이 관측 파이프라인을 거쳐 대시보드와 Jaeger에 성공적으로 기록되는지 확인하는 데 목적이 있다. "
        "통제된 환경에서 확인된 시나리오별 관측 사실(Fact)은 다음과 같다.",
    )
    add_paragraph(
        doc,
        "normal (정상 흐름): 요청 발생 시 고유 request_id와 trace_id가 생성되며, 대시보드에 상태(Status OK)와 메타데이터가 정상적으로 표출됨을 확인하였다.",
        bold_prefix="normal (정상 흐름):",
    )
    add_paragraph(
        doc,
        "slow_retrieve (RAG 검색 지연): 의도적인 검색 지연 주입 시, 전체 응답 시간 대비 retrieve span의 소요 시간이 급증하여 병목 구간으로 기록됨을 확인하였다.",
        bold_prefix="slow_retrieve (RAG 검색 지연):",
    )
    add_image(doc, "08_chat_slow_retrieve_response.png", "그림 14. RAG/검색 지연 시나리오: retrieve와 llm_call 중 병목 후보를 설명하는 채팅 화면")
    add_paragraph(
        doc,
        "slow_llm (모델 생성 지연): 모델 추론 지연 주입 시, 3단계 span 중 llm_call span의 소요 시간이 가장 긴 구간으로 고립되어 기록됨을 확인하였다.",
        bold_prefix="slow_llm (모델 생성 지연):",
    )
    add_image(doc, "09_chat_slow_llm_response.png", "그림 15. 모델 응답 지연 시나리오: llm_call 지연과 모델 실행 리소스 점검 포인트")
    add_paragraph(
        doc,
        "high_token (토큰 과다 사용): 긴 프롬프트 입력 시, Prometheus에 수집된 Token Usage 지표와 추정 비용(Cost) 지표가 동반 상승함을 확인하였다.",
        bold_prefix="high_token (토큰 과다 사용):",
    )
    add_image(doc, "10_chat_high_token_response.png", "그림 16. 토큰 과다 사용 시나리오: 토큰 사용량과 비용 증가 관계를 설명하는 화면")
    add_paragraph(
        doc,
        "error (실패 요청): 인위적인 실패 요청 발생 시, HTTP 에러로 소실되지 않고 status=error 상태의 메타데이터와 trace_id가 유지되어 대시보드 최근 요청 목록에 기록됨을 확인하였다.",
        bold_prefix="error (실패 요청):",
    )
    add_image(doc, "11_chat_error_response.png", "그림 17. 실패 요청 시나리오: ERROR 상태와 request_id, trace_id 기반 추적 정보")
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "수집된 관측 데이터가 가지는 공학적 의미와 시스템 병목 원인에 대한 상세 추론은 8.1절에서 논의한다.",
    )
    doc.add_heading("7.3 검증 흐름과 판정 기준", level=2)
    add_paragraph(
        doc,
        "본 검증의 최종 판정 기준은 단일 요청에 대한 End-to-End 추적 가능성 확보 여부에 있다. Chat UI에서 생성된 요청의 request_id와 trace_id가 대시보드의 최근 요청 목록에 표시되고, "
        "Jaeger에서는 같은 trace 안에서 retrieve, llm_call, postprocess span이 확인되어야 한다. Prometheus와 Grafana에서는 같은 기간의 latency, token, cost, error 지표가 "
        "시나리오의 의도와 일치하는 방향으로 변해야 한다.",
    )
    add_paragraph(
        doc,
        "이 기준을 충족하면 OLLY는 운영자가 시스템 지연 및 비용 증가 원인과 실패 발생 지점을 데이터에 근거하여 추적할 수 있는 관측성 플랫폼임을 입증한다.",
    )
    add_image(doc, "05_grafana.png", "그림 18. Grafana 연계 대시보드 접속 화면: Prometheus 기반 시각화 도구")

    doc.add_heading("8. 논의 및 향후 확장 방향", level=1)
    add_paragraph(
        doc,
        "본 장은 검증 과정에서 수집한 관측 데이터를 기반으로 OLLY의 운영적 의미와 확장 가능성을 논의한다. "
        "본 시스템은 비용, 지연, 토큰 사용량, 실패 내역 및 병목 구간 등을 집중적으로 추적하는 운영 관측성(Operational Observability) 특화 플랫폼으로 설계되었다.",
    )
    add_paragraph(
        doc,
        "본 논의는 하나의 LLM 요청이 request_id와 trace_id를 기준으로 trace, metric, dashboard, alert까지 연결되는 구조와 "
        "운영자가 해당 데이터를 근거로 원인과 대응 방향을 판단하는 과정을 중심으로 구성된다.",
    )
    doc.add_heading("8.1 관측 시나리오 기반의 시스템 병목 데이터 분석", level=2)
    add_paragraph(
        doc,
        "본 절에서는 앞서 구축한 OLLY 관측 파이프라인을 통해 5가지 통제된 시나리오(normal, slow_retrieve, slow_llm, high_token, error)를 실행하고, "
        "수집된 데이터를 바탕으로 시스템의 병목 구간과 인프라의 이상 징후를 공학적으로 추론(Deduction)한다.",
    )
    add_paragraph(
        doc,
        "특히 분산 환경 모니터링 아키텍처의 실효성을 엄밀하게 검증하기 위해, 각 시나리오별로 10회의 연속 부하(Load)를 발생시켜 총 50회의 요청을 테스트하였다. "
        "이 과정에서 OpenTelemetry 파이프라인의 컨텍스트 유실은 단 한 건도 발생하지 않았으며, 수집된 데이터를 바탕으로 산출된 SRE 핵심 지표(지연 시간 중앙값, 에러율, 토큰 사용량)는 표 1과 같다.",
    )
    add_paragraph(
        doc,
        "시스템 관측성(Observability) 관점에서 산술 평균(Average)은 간헐적인 네트워크 스파이크 등의 아웃라이어에 의해 왜곡되기 쉬우므로, "
        "일반적인 사용자 경험을 대변하는 p50(Median)과 상위 5%의 지연을 보여주는 p95 지표를 기준으로 데이터를 집계하였다.",
    )
    add_caption(doc, "표 1. 시나리오별 10회 반복 부하 검증 관측 지표 (p50/p95 기준)")
    add_table(
        doc,
        ["검증 시나리오", "에러율", "p50 지연 (Median)", "p95 지연 (Tail)", "총 토큰 (Avg)", "이상 감지 구간 (Anomaly Span, p50 기준)"],
        [
            ["normal (정상 흐름)", "0%", "1.47초", "3.02초", "243개", "llm_call (1.23초)"],
            [
                "slow_retrieve (RAG 지연)",
                "0%",
                "3.96초",
                "5.93초",
                "270개",
                "retrieve (1.80초, 정상 대비 약 12배 폭증)",
            ],
            ["slow_llm (생성 지연)", "0%", "2.23초", "3.11초", "271개", "llm_call (1.99초)"],
            [
                "high_token (토큰 과다)",
                "0%",
                "34.58초",
                "36.86초",
                "457개",
                "llm_call (34.35초, 리소스 한계 도달)",
            ],
            ["error (실패 요청)", "100%", "0.15초", "0.15초", "0개", "retrieve (HTTP Error 조기 실패)"],
        ],
        [1.0, 0.55, 0.78, 0.78, 0.72, 2.9],
    )
    add_paragraph(
        doc,
        "위와 같이 파이프라인이 안정적으로 집계해 낸 전체 통계를 바탕으로, 아래에서는 각 시나리오별 대표 Trace를 고립(Isolation)시켜 세부적인 원인을 분석한다. "
        "단순한 대시보드 기능 설명을 지양하고, 관측 데이터가 가리키는 실제 시스템의 물리적 상태와 운영 인사이트를 도출하는 데 집중하였다.",
    )
    add_paragraph(
        doc,
        "RAG 검색 병목(slow_retrieve) 데이터 해석 및 고립화(Isolation): slow_retrieve 시나리오는 p50 지연이 3.96초로 상승했으며, retrieve span이 1.80초로 정상 흐름 대비 약 12배 증가하였다. "
        "이는 모델 추론 성능만으로는 설명할 수 없는 검색 계층 병목이 존재함을 의미한다. 운영자는 이 수치를 근거로 모델 교체나 GPU 증설보다 인덱스 최적화, top-k 조정, 검색 캐시 도입, chunk 전략 개선과 같은 검색 계층 중심의 튜닝을 우선 검토할 수 있다.",
        bold_prefix="RAG 검색 병목(slow_retrieve) 데이터 해석 및 고립화(Isolation):",
    )
    add_paragraph(
        doc,
        "모델 추론 지연(slow_llm)과 로컬 리소스 한계의 상관관계: slow_llm 시나리오는 p50 지연 2.23초, llm_call span 1.99초로 관측되었다. "
        "검색과 후처리 단계가 상대적으로 짧게 유지되는 상황에서 llm_call 구간이 지연의 대부분을 차지하므로, 해당 병목은 로컬 SLM(Ollama gemma3:1b)의 모델 생성 경로와 직접적으로 연결된다. "
        "이는 OLLY의 다단계 span 구조가 CPU 기반 로컬 추론 환경에서 발생하는 generation latency spike를 llm_call 컨텍스트로 고립시켜 시각화할 수 있음을 보여준다.",
        bold_prefix="모델 추론 지연(slow_llm)과 로컬 리소스 한계의 상관관계:",
    )
    add_paragraph(
        doc,
        "토큰 스파이크(high_token)에 따른 비용 변동성 가시화: high_token 시나리오는 평균 총 토큰 수가 457개까지 증가했고, p50 지연은 34.58초, llm_call span은 34.35초로 관측되었다. "
        "현재 MVP는 로컬 SLM을 사용하므로 gemma3:1b의 token_cost_usd는 0으로 계산되지만, 긴 문맥은 로컬 추론 시간을 증가시켜 추정 infra_cost_usd와 지연을 함께 확대한다. "
        "따라서 이 시나리오는 과도한 context injection이나 비효율적인 prompt template이 운영 비용 추정치와 성능 저하로 이어질 수 있음을 정량적으로 드러낸다.",
        bold_prefix="토큰 스파이크(high_token)에 따른 비용 변동성 가시화:",
    )
    add_paragraph(
        doc,
        "실패 요청(error)의 추적 가능성: error 시나리오는 10회 모두 status=error로 기록되었고, p50과 p95 지연은 모두 0.15초로 집계되었다. "
        "실패 요청이 HTTP 오류로 소실되지 않고 request_id와 trace_id를 유지했다는 점은 장애 분석을 위한 관측 단위가 보존되었음을 의미한다. "
        "운영자는 해당 trace를 기준으로 실패 단계, 최근 실패율, 알림 조건을 연계하여 장애 대응 우선순위를 결정할 수 있다.",
        bold_prefix="실패 요청(error)의 추적 가능성:",
    )
    add_paragraph(
        doc,
        "그림 19~23은 위와 같은 공학적 데이터 해석을 뒷받침하는 OLLY 플랫폼의 실제 화면 근거이다. 운영자는 운영 요약 화면에서 전체 실패율과 비용 신호를 확인하고, 요청 추적 상세 화면에서 병목 span을 고립한 뒤, "
        "실시간 지표 화면에서 Prometheus/Grafana 지표가 동일한 방향성을 보이는지 교차 검증할 수 있다.",
    )
    add_image(
        doc,
        "19_dashboard_scenario_overview.png",
        "그림 19. 대시보드 검증 종합 화면: 실패율, 평균 지연, 토큰, 비용과 최근 시나리오 요청 확인",
        6.8,
    )
    add_image(
        doc,
        "20_dashboard_trace_slow_retrieve.png",
        "그림 20. slow_retrieve 요청 추적 상세 화면: retrieve span이 가장 긴 병목으로 표시되는 화면",
        6.8,
    )
    add_image(
        doc,
        "21_dashboard_trace_slow_llm.png",
        "그림 21. slow_llm 요청 추적 상세 화면: LLM Generation span이 주요 병목으로 표시되는 화면",
        6.8,
    )
    add_image(
        doc,
        "22_dashboard_trace_error.png",
        "그림 22. error 요청 추적 상세 화면: 실패 요청이 ERROR 상태와 request_id 기준으로 추적되는 화면",
        6.8,
    )
    add_image(
        doc,
        "23_dashboard_signals_metrics.png",
        "그림 23. 실시간 지표 화면: Grafana 패널에서 latency, token, cost, stage p95를 확인",
        6.8,
    )
    add_table(
        doc,
        ["시나리오", "관측 근거", "실무 해석", "실무자 조치"],
        [
            [
                "slow_retrieve",
                "retrieve span 또는 stage duration 증가",
                "RAG 검색, Vector DB 조회, 문서 검색 단계가 병목일 가능성",
                "인덱스 최적화, top-k 조정, 검색 캐시, chunk 전략 개선",
            ],
            [
                "slow_llm",
                "llm_call span과 전체 latency 증가",
                "모델 생성 또는 외부 LLM API 응답 지연이 전체 요청을 지배",
                "max token 제한, streaming 적용, 모델 교체, GPU/리소스 증설, fallback model 검토",
            ],
            [
                "high_token",
                "input/output token과 cost 지표 증가",
                "프롬프트 템플릿 또는 검색 문맥이 과도해 비용과 지연을 키움",
                "prompt template 축소, context 제한, 요약/압축, 기능별 비용 budget 설정",
            ],
            [
                "error",
                "error status, error metric, 실패 trace 기록",
                "실패 요청이 사라지지 않고 특정 request_id와 trace_id로 추적 가능",
                "예외 유형 분류, retry 정책, alert rule 조정, 장애 runbook 작성",
            ],
        ],
        [1.05, 1.9, 1.85, 1.9],
    )
    add_paragraph(
        doc,
        "각 시나리오는 실제 운영 질문을 대표한다. slow_retrieve와 slow_llm은 응답 지연 현상을 검색 병목과 모델 병목으로 분리하고, "
        "high_token은 비용 증가가 어느 요청 패턴에서 시작되는지 제시하며, error는 실패 요청을 request 단위로 기록하여 장애 대응을 위한 식별 기준을 제공한다.",
    )
    doc.add_heading("8.2 CNCF 기반 관측성 적용의 의미", level=2)
    add_paragraph(
        doc,
        "CNCF 기반 관측성 구성에서 Kubernetes는 sample-llm-api, Ollama, OpenTelemetry Collector, Prometheus, Jaeger, Grafana를 "
        "재현 가능한 컨테이너 실행 환경으로 구성하였고, OpenTelemetry는 /chat 요청 내부의 retrieve, llm_call, postprocess 단계를 span으로 분해하였다.",
    )
    add_paragraph(
        doc,
        "Prometheus는 요청 수, latency, token, cost, error, stage duration을 운영 지표와 alert rule의 근거로 활용하며, Jaeger는 특정 request_id와 trace_id의 병목을 시각적으로 확인하도록 지원한다. "
        "즉 OLLY는 LLM 운영 문제를 CNCF 관측성 구조로 바꾸어 설명한 사례이다.",
    )
    doc.add_heading("8.3 실무 적용 가능성과 전제 조건", level=2)
    add_paragraph(
        doc,
        "현재 구현은 로컬 SLM 기반 MVP이지만, 구조적으로는 기존 사용자용 LLM/RAG 서비스의 /chat API에 계측을 추가하는 방식으로 확장할 수 있다. "
        "서비스 로직을 유지하면서 요청 처리 경로에 trace와 metric 기록을 추가하고, 운영자가 같은 request_id와 trace_id로 대시보드와 Jaeger를 오가며 확인하는 구조이다.",
    )
    add_paragraph(
        doc,
        "실무 적용을 위해서는 raw prompt 저장 정책, 개인정보 마스킹, 장기 metric 저장, 권한 관리, SLO 기준, 비용 산정 기준이 추가로 필요하다. 특히 raw prompt를 그대로 저장하면 보안 위험이 커지므로 "
        "prompt_template_id, prompt_length, token count, 익명화된 user/org id처럼 운영 분석에 필요한 최소 메타데이터부터 남기는 방식이 적절하다.",
    )
    doc.add_heading("8.4 시스템 성능 중심의 관측성 아키텍처가 가지는 의의", level=2)
    add_paragraph(
        doc,
        "현재 OLLY가 다루는 핵심 영역은 LLM 애플리케이션의 시스템 운영 관측성이다. 실무적인 LLM 서비스 아키텍처에서 답변의 품질(Semantic Quality)과 시스템의 성능(Operational Performance)은 구분되어 모니터링되어야 한다. "
        "본 MVP는 이 중 인프라 병목 구간 식별과 리소스 비용 통제라는 시스템적 문제를 해결하는 데 특화되어 있다.",
    )
    add_paragraph(
        doc,
        "향후 RAG 근거 일치도 검사, LLM-as-a-judge 기반 자동 평가와 같은 품질 모니터링 기능이 요구되더라도, 현재 OLLY가 구축한 request_id 및 trace_id 기반의 분산 추적 아키텍처 위에서 해당 평가 지표들을 새로운 메타데이터로 통합할 수 있다. "
        "즉 OLLY는 모니터링 화면을 넘어 향후 품질 평가 모델을 수용할 수 있는 견고한 관측성 기반(Observability Foundation)을 마련했다는 데 구조적 의의가 있다.",
    )
    doc.add_heading("8.5 프롬프트 관측의 보안 기준", level=2)
    add_paragraph(
        doc,
        "프롬프트 관측 시 식별자(request_id, trace_id), 비즈니스 컨텍스트(feature, prompt_template_id), 토큰 및 비용 지표(input/output_tokens, cost_usd), "
        "그리고 성능 지표(latency_ms, stage durations)를 결합한 통합 메타데이터 스키마를 관측 단위로 정의하였다.",
    )
    add_paragraph(
        doc,
        "프롬프트 관측에는 개인정보와 내부 정보 유출 위험이 수반된다. 따라서 실무형 OLLY에서는 raw prompt 전체 저장보다 계측 목적에 필요한 메타데이터를 먼저 정의하고, "
        "원문이 필요한 경우에는 샘플링, 마스킹, 보존 기간, 접근 권한 정책을 함께 적용해야 한다.",
    )
    doc.add_heading("8.6 분산 환경 아키텍처 구현 및 데이터 정합성 트러블슈팅", level=2)
    add_paragraph(
        doc,
        "시스템 파이프라인 구축 과정에서는 분산 컨테이너 간 네트워크 라우팅, OTLP 통신 프로토콜 설정, 비동기 요청 처리 경로에서의 trace 컨텍스트 유지, "
        "Prometheus metric과 Jaeger trace 간 데이터 정합성 확보와 같은 시스템 소프트웨어 레벨의 기술적 과제가 발생하였다. "
        "본 프로젝트에서는 오픈소스 도구를 단순 실행하는 수준에 머무르지 않고, 각 컴포넌트의 포트, 프로토콜, 식별자 전파 경로를 명시적으로 제어하여 관측 파이프라인을 구성하였다.",
    )
    doc.add_heading("8.6.1 OTLP 통신 프로토콜 정합성 및 대시보드 데이터 연계", level=3)
    add_paragraph(
        doc,
        "이슈: sample-llm-api에서 생성한 계측 데이터를 OpenTelemetry Collector로 전송하는 과정에서 OTLP endpoint와 Kubernetes Service 포트가 일관되지 않으면 trace 수집이 중단될 수 있었다. "
        "또한 metric과 trace가 각각 수집되더라도 Prometheus 시계열 지표(Metric)와 Jaeger 분산 추적(Trace)이 동일한 요청 식별자로 연결되지 않으면, 특정 latency spike가 어느 요청에서 발생했는지 역추적하기 어려운 구조적 단절이 발생한다.",
        bold_prefix="이슈:",
    )
    add_paragraph(
        doc,
        "기술적 해결 과정: OTel Collector의 OTLP receiver는 gRPC 4317과 HTTP 4318을 모두 열도록 구성하고, sample-llm-api의 OTLP exporter는 gRPC exporter인 OTLPSpanExporter를 사용하여 "
        "OTEL_EXPORTER_OTLP_ENDPOINT=http://otel-collector:4317로 고정하였다. Kubernetes manifest에서도 otel-collector Service의 port와 targetPort를 4317, 4318로 각각 맞추고, "
        "Jaeger는 COLLECTOR_OTLP_ENABLED=true와 jaeger:4317 경로로 trace를 수신하도록 구성하였다. 대시보드 접근 경로는 sample-llm-api NodePort 30001, Jaeger UI NodePort 30002로 분리하여 "
        "관측 대상 API와 trace 분석 화면의 네트워크 경로를 명확히 구분하였다.",
        bold_prefix="기술적 해결 과정:",
    )
    add_paragraph(
        doc,
        "데이터 정합성은 FastAPIInstrumentor가 생성하는 ASGI 계측 계층과 /chat 핸들러 내부의 명시적 식별자 주입 로직을 결합하여 확보하였다. 요청 처리 시작 시 request_id를 UUID 기반으로 생성하고, "
        "OpenTelemetry의 current root span에서 trace_id를 추출한 뒤 olly.request_id, olly.trace_id, olly.feature, olly.scenario 속성을 span attribute로 기록하였다. "
        "동시에 Prometheus의 olly_request_info_total에는 request_id, trace_id, model, feature, scenario, status를 label로 기록하였다. "
        "이 구조를 통해 대시보드의 Recent Requests, Prometheus 지표, Jaeger trace detail이 동일한 request_id와 trace_id를 기준으로 1:1 연계된다.",
    )
    doc.add_heading("8.6.2 비동기 처리 환경의 Span 구조화 및 알림 파이프라인 제어", level=3)
    add_paragraph(
        doc,
        "이슈: 전체 API 응답 시간만 측정하는 방식으로는 RAG 검색 지연, LLM 모델 추론 지연, 후처리 지연을 구분할 수 없었다. 특히 FastAPI의 async/await 기반 코루틴 흐름에서는 retrieve, llm_call, postprocess가 모두 비동기 함수로 실행되므로, "
        "각 await 경계에서 상위 trace와 하위 span의 관계가 유지되어야 Jaeger에서 단일 Trace Tree로 병목 구간을 해석할 수 있다.",
        bold_prefix="이슈:",
    )
    add_paragraph(
        doc,
        "기술적 해결 과정: /chat 핸들러 내부에서 OpenTelemetry의 tracer.start_as_current_span()을 사용해 retrieve, llm_call, postprocess를 하위 span으로 명시적으로 분리하였다. "
        "각 span은 await retrieve(), await llm_client.llm_call(), await postprocess() 호출을 직접 감싸도록 배치하여 비동기 호출 이후에도 동일한 trace context 안에서 단계별 duration이 기록되도록 하였다. "
        "또한 Prometheus에는 stage_timer()를 통해 olly_stage_duration_seconds{stage=\"retrieve\"}, olly_stage_duration_seconds{stage=\"llm_call\"}, olly_stage_duration_seconds{stage=\"postprocess\"} histogram을 기록하였다. "
        "그 결과 Jaeger에서는 단일 trace 하위에 3단계 span이 계층적으로 표시되고, Prometheus에서는 동일 단계를 시계열 지표로 집계할 수 있다.",
        bold_prefix="기술적 해결 과정:",
    )
    add_paragraph(
        doc,
        "알림 파이프라인 제어: sample-llm-api 내부에 비동기 백그라운드 task 기반 AlertEvaluator를 구현하였다. "
        "AlertEvaluator는 ALERT_EVAL_INTERVAL_SECONDS=30 주기로 Prometheus HTTP API를 조회하고, p95 latency, error rate, token rate, cost rate, retrieve p95와 같은 PromQL 조건을 평가한다. "
        "규칙별 cooldown_seconds를 적용하여 순간적인 부하로 인한 알림 폭주를 방지하고, 조건이 충족된 경우 httpx.AsyncClient로 Discord Webhook에 알림을 전송한다. "
        "이로써 운영자는 알림 수신 후 동일 시간대의 request_id와 trace_id를 대시보드와 Jaeger에서 추적하는 대응 흐름을 수행할 수 있다.",
        bold_prefix="알림 파이프라인 제어:",
    )
    doc.add_heading("8.7 AI 활용 및 아키텍처 설계 주도권 통제", level=2)
    add_paragraph(
        doc,
        "본 프로젝트에서 생성형 AI는 아이디어 정리, MVP 범위 설정, FastAPI 코드 뼈대 및 Kubernetes 설정 파일 작성 등 단순 반복 작업을 가속하는 도구로 활용되었다. "
        "특히 LLM 서비스 운영자가 필요로 하는 관측 정보를 구체화하는 기획 초기 단계에서 유용했다. "
        "그러나 최종 아키텍처 방향성과 핵심 설계는 AI가 생성한 초안을 비판적으로 검토하고 통제한 팀의 주도적인 공학적 판단에 따라 완성되었다.",
    )
    doc.add_heading("8.7.1 외부 LLM API 중심 구조 배제 및 로컬 SLM 도입", level=3)
    add_paragraph(
        doc,
        "초기 AI 생성 아키텍처 초안: 시스템 설계 초기 단계에서 AI 도구는 개발 편의성과 범용성을 이유로 OpenAI, Claude 등 외부 클라우드 LLM API를 연동하는 파이프라인 초안을 주로 생성하였다.",
        bold_prefix="초기 AI 생성 아키텍처 초안:",
    )
    add_paragraph(
        doc,
        "팀의 공학적 설계 통제: 팀 내부 논의 과정에서, 과거 Agentic AI 프로젝트를 통해 다수의 상용 LLM과 로컬 SLM을 직접 구축 및 운영해 본 팀원의 실무적 경험과 인사이트를 핵심 판단 근거로 삼았다. "
        "외부 API를 맹목적으로 수용할 경우 학내 실험 및 검증 환경의 네트워크 불안정성, API 호출 쿼터 제한(Rate Limit), 예기치 못한 비용 폭증 등으로 인해 반복 검증의 안정성이 치명적으로 훼손된다는 실증적 한계를 사전에 인지하였다. "
        "이에 따라 AI가 제안한 외부 의존적 구조를 전면 기각하고, 인프라 비용 추정이 명확하며 팀원 모두가 동일한 로컬 컴퓨팅 리소스 상에서 통제된 반복 테스트를 수행할 수 있는 로컬 SLM(Ollama gemma3:1b) 기반 인프라로 아키텍처를 주도적으로 재설계하였다.",
        bold_prefix="팀의 공학적 설계 통제:",
    )
    doc.add_heading("8.7.2 핵심 가치 집중을 위한 기능 스코핑(Scoping) 제어", level=3)
    add_paragraph(
        doc,
        "초기 AI 생성 아키텍처 초안: AI는 관측 플랫폼 구축 시 로그인/회원가입, 관리자 권한 관리(RBAC), 멀티테넌시, 과금 시스템 등을 포함한 방대한 상용 웹 서비스 수준의 기능 초안을 제안하였다.",
        bold_prefix="초기 AI 생성 아키텍처 초안:",
    )
    add_paragraph(
        doc,
        "팀의 공학적 설계 통제: 본 프로젝트의 본질적 가치는 일반적인 웹 서비스 구현이 아니라 거대 모델 요청 처리에 특화된 운영 관측성(Operational Observability) 확보에 있다. "
        "따라서 팀은 부차적인 관리 기능들을 MVP 구현 범위에서 과감히 잘라내는(Scoping) 하향식 제어를 수행하였다. "
        "불필요한 마이너 기능 구현에 소모될 리소스를 차단하고, 확보된 자원을 request_id 및 trace_id 기반의 지연 시간, 비용, 병목 분석 파이프라인의 완성도를 높이는 데 집중 투입하였다.",
        bold_prefix="팀의 공학적 설계 통제:",
    )
    add_paragraph(
        doc,
        "결과적으로 본 프로젝트는 AI가 생성한 초안을 단순 수용하는 데 그치지 않고, 명확한 관측성 목표와 현실적 제약을 바탕으로 시스템 범위를 튜닝하여 "
        "CNCF 도구 기반의 통합 플랫폼 MVP를 독자적으로 완성해 냈다는 데 의의가 있다.",
    )
    doc.add_heading("8.8 향후 확장 방향", level=2)
    add_paragraph(
        doc,
        "본 MVP의 한계를 보완하고 실무 프로덕션 환경에서의 적용성을 극대화하기 위해, 향후 확장은 품질 평가 영역과 엔터프라이즈 운영 인프라 영역으로 구분하여 추진할 수 있다.",
    )
    doc.add_heading("8.8.1 품질 평가(Semantic Monitoring) 영역 확장", level=3)
    add_bullets(
        doc,
        [
            "RAG 근거 문서와 답변 문장의 일치도 검사 및 인용(Citation) 검증",
            "LLM-as-a-judge 기반 자동 평가 및 사용자 피드백/운영자 라벨링 데이터 수집 체계 구축",
            "정답 데이터셋 기반 Regression Eval 파이프라인 구성",
        ],
    )
    doc.add_heading("8.8.2 엔터프라이즈 운영 인프라 고도화", level=3)
    add_bullets(
        doc,
        [
            "실제 상용 Vector DB, 인증/권한 체계, 장기 로그 저장소 연동",
            "운영 환경용 Ingress/TLS 적용, Secret 관리, RBAC 및 감사 로그 체계 도입",
            "다중 조직별 비용 분리 및 SLO 기반 커스텀 Alert Rule 확장",
        ],
    )

    doc.add_heading("9. 결론", level=1)
    add_paragraph(
        doc,
        "OLLY는 LLM 서비스 운영자가 실제 프로덕션 환경에서 직면하는 비용 폭증, 응답 지연, 토큰 소모 스파이크 및 시스템 실패 원인을 request_id와 trace_id를 통해 하부 인프라부터 응답 단계까지 End-to-End로 정밀 추적하도록 지원하는 실무형 관측성 플랫폼이다. "
        "요청마다 Prometheus metric과 Jaeger trace를 결합해 운영 질문에 대응하도록 구현하였다.",
    )
    add_paragraph(
        doc,
        "본 프로젝트의 주요 의의는 CNCF 생태계의 표준 도구들을 유기적으로 결합(Kubernetes 기반 인프라, OpenTelemetry 계측, Prometheus 지표/알림, Jaeger 분산 추적)하여 "
        "LLM 서비스 운영 환경에서의 실효성을 성공적으로 입증한 것에 있다.",
    )
    add_paragraph(
        doc,
        "또한 OLLY는 실무 적용을 전제로 기존 서비스에 연결되는 관측성 계층으로 설계되었다. 운영자는 request_id와 trace_id를 기준으로 "
        "사용자 요청, 비용, 토큰, 단계별 병목, 알림을 연결해서 확인할 수 있으며, 이는 실제 LLM 서비스 운영에서 필요한 장애 대응과 "
        "비용 통제 절차에 직접 대응한다.",
    )
    add_paragraph(
        doc,
        "본 프로젝트의 최종 산출물은 시스템 비용, 재현성, 실제 운영 프로세스 등의 현실적 제약을 구조적으로 해결한 통합 아키텍처이다. "
        "관측 대상 정의, 로컬 SLM 선택, request_id와 trace_id 중심의 관측 단위, 검증 시나리오, Discord 알림 방식은 프로젝트 목표와 구현 제약을 기준으로 확정하였다.",
    )
    add_paragraph(
        doc,
        "본 프로젝트는 관측 대상을 로컬 SLM 환경으로 통제하여 실험의 재현성을 확보하고, CNCF 표준 도구들을 유기적으로 결합하여 통합 모니터링 파이프라인을 완성하였다. "
        "본 MVP를 통해 확립된 Trace-Metric 연계 아키텍처는 향후 대규모 외부 LLM API 서비스나 정교한 RAG 챗봇 시스템으로 유연하게 이식(Porting)될 수 있는 견고한 시스템 기반을 제공한다.",
    )

    doc.add_page_break()
    doc.add_heading("부록 A. 주요 PromQL", level=1)
    add_code_block(
        doc,
        """
sum(increase(olly_requests_total[1h]))
sum(increase(olly_tokens_total[1h])) by (feature)
sum(increase(olly_cost_usd_total[1h])) by (feature)
histogram_quantile(0.95, sum by (stage, le) (increase(olly_stage_duration_seconds_bucket[1h])))
sum(increase(olly_errors_total[5m])) / clamp_min(sum(increase(olly_requests_total[5m])), 1)
        """,
    )
    doc.add_heading("부록 B. 주요 Metric", level=1)
    add_table(
        doc,
        ["Metric", "Type", "의미"],
        [
            ["olly_requests_total", "Counter", "LLM chat 요청 수"],
            ["olly_tokens_total", "Counter", "입력/출력 토큰 사용량"],
            ["olly_cost_usd_total", "Counter", "토큰 비용과 로컬 인프라 비용을 합산한 누적 추정 비용"],
            ["olly_token_cost_usd_total", "Counter", "외부 API형 모델에서의 토큰 기반 누적 비용"],
            ["olly_infra_cost_usd_total", "Counter", "로컬 모델 실행 시간 기반 누적 인프라 비용"],
            ["olly_stage_duration_seconds", "Histogram", "retrieve, llm_call, postprocess 단계별 소요 시간 분포"],
            ["olly_request_duration_seconds", "Histogram", "요청 전체 소요 시간 분포"],
            ["olly_errors_total", "Counter", "실패 요청 수"],
        ],
        [2.2, 1.1, 3.0],
    )
    doc.add_heading("부록 C. Kubernetes 검증 출력", level=1)
    add_paragraph(
        doc,
        "본문 6장에서는 Kubernetes 검증 결과를 요약하고, 전체 kubectl 출력은 재현성 확인을 위해 부록에 수록한다.",
    )
    add_code_block(doc, k8s_status)

    doc.core_properties.title = "OLLY 최종 보고서"
    doc.core_properties.subject = "LLM 서비스 관측성 플랫폼 MVP"
    doc.core_properties.author = "OLLY 5조"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
